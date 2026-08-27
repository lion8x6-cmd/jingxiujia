#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for sec in doc.sections:
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

for lv, sz in [(1,18),(2,14),(3,12)]:
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = '微软雅黑'
    hs.font.size = Pt(sz)
    hs.font.bold = True
    hs.font.color.rgb = [None, RGBColor(0x1A,0x3C,0x6E), RGBColor(0x2C,0x5F,0xA0), RGBColor(0x3A,0x7A,0xC4)][lv]
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.paragraph_format.space_before = Pt(18 if lv==1 else 12)
    hs.paragraph_format.space_after = Pt(8)

def TBL(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text=''
        r = c.paragraphs[0].add_run(h)
        r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        r.font.name='微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C5FA0"/>'))
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text=''
            r = c.paragraphs[0].add_run(str(val))
            r.font.size=Pt(9.5); r.font.name='微软雅黑'
            r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
            if ri%2==1:
                c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F0FE"/>'))
    doc.add_paragraph()

def IMG(cap):
    pp = doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = pp.add_run(f'【{cap}】')
    r.font.size=Pt(10); r.font.color.rgb=RGBColor(0x99,0x99,0x99)
    r.font.italic=True; r.font.name='微软雅黑'
    r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    pPr=pp._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/><w:left w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/><w:bottom w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/><w:right w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/></w:pBdr>'))
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>'))
    pp.paragraph_format.space_before=Pt(8); pp.paragraph_format.space_after=Pt(8)

def B(text):
    pp=doc.add_paragraph(style='List Bullet'); pp.clear()
    r=pp.add_run(text); r.font.size=Pt(10.5); r.font.name='微软雅黑'
    r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')

def P(text, bold=False):
    pp=doc.add_paragraph()
    pp.paragraph_format.first_line_indent=Cm(0.74)
    r=pp.add_run(text); r.font.size=Pt(10.5); r.font.name='微软雅黑'; r.bold=bold
    r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')

def RP(segs):
    pp=doc.add_paragraph()
    pp.paragraph_format.first_line_indent=Cm(0.74)
    for txt,bold in segs:
        r=pp.add_run(txt); r.font.size=Pt(10.5); r.font.name='微软雅黑'; r.bold=bold
        r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')

def PB():
    doc.add_page_break()

# ========== 封面 ==========
for _ in range(6): doc.add_paragraph()
for txt,sz,col in [('AI一键精修',32,RGBColor(0x1A,0x3C,0x6E)),('市场需求文档（MRD）',22,RGBColor(0x2C,0x5F,0xA0))]:
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=pp.add_run(txt); r.font.size=Pt(sz); r.font.bold=True; r.font.color.rgb=col
    r.font.name='微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
doc.add_paragraph()
pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=pp.add_run('Market Requirements Document'); r.font.size=Pt(14)
r.font.color.rgb=RGBColor(0x88,0x88,0x88); r.font.italic=True
for _ in range(4): doc.add_paragraph()
it=doc.add_table(rows=5,cols=2); it.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate([('产品名称','AI一键精修'),('文档版本','V1.0'),('编写日期','2026年8月14日'),('文档状态','初稿'),('密级','内部保密')]):
    for j,val in enumerate([k,v]):
        c=it.rows[i].cells[j]; c.text=''
        r=c.paragraphs[0].add_run(val); r.font.size=Pt(11); r.font.name='微软雅黑'
        r.bold=(j==0); r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
        if j==0: r.font.color.rgb=RGBColor(0x2C,0x5F,0xA0)
PB()

# ========== 目录 ==========
doc.add_heading('目  录',level=1)
for item in ['一、市场现状与行业格局','二、用户画像与核心场景','三、市场验证（抖音实测）','四、竞品分析','五、产品差异化优势','六、产品功能规划','七、商业变现','八、后续规划','九、风险与应对','附录']:
    pp=doc.add_paragraph(); r=pp.add_run(item); r.font.size=Pt(12); r.font.name='微软雅黑'
    r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); pp.paragraph_format.space_after=Pt(8)
PB()

exec(open('/Users/mr.baihe/WorkBuddy/P 图修图-精修家/output/mrd-20260814/working/content_part1.py').read())
exec(open('/Users/mr.baihe/WorkBuddy/P 图修图-精修家/output/mrd-20260814/working/content_part2.py').read())

out = '/Users/mr.baihe/WorkBuddy/P 图修图-精修家/output/mrd-20260814/stage3/AI一键精修_MRD_V1.0.docx'
doc.save(out)
print(f'SAVED: {out}')
