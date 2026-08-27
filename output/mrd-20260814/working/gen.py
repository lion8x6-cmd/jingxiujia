#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()
s = doc.styles['Normal']
s.font.name='微软雅黑'; s.font.size=Pt(10.5); s.font.color.rgb=RGBColor(0x33,0x33,0x33)
s.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.5
for sec in doc.sections:
    sec.top_margin=Cm(2.54); sec.bottom_margin=Cm(2.54); sec.left_margin=Cm(3.17); sec.right_margin=Cm(3.17)
for lv,sz in [(1,18),(2,14),(3,12)]:
    h=doc.styles[f'Heading {lv}']; h.font.name='微软雅黑'; h.font.size=Pt(sz); h.font.bold=True
    h.font.color.rgb=[None,RGBColor(0x1A,0x3C,0x6E),RGBColor(0x2C,0x5F,0xA0),RGBColor(0x3A,0x7A,0xC4)][lv]
    h.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    h.paragraph_format.space_before=Pt(18 if lv==1 else 12); h.paragraph_format.space_after=Pt(8)

def TBL(hd,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(hd)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hd):
        c=t.rows[0].cells[i]; c.text=''; r=c.paragraphs[0].add_run(h)
        r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.font.name='微软雅黑'
        r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C5FA0"/>'))
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; r=c.paragraphs[0].add_run(str(val))
            r.font.size=Pt(9.5); r.font.name='微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
            if ri%2==1: c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F0FE"/>'))
    doc.add_paragraph()

def IMG(cap):
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=pp.add_run(f'\u3010{cap}\u3011'); r.font.size=Pt(10); r.font.color.rgb=RGBColor(0x99,0x99,0x99)
    r.font.italic=True; r.font.name='微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    pPr=pp._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/><w:left w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/><w:bottom w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/><w:right w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/></w:pBdr>'))
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>'))
    pp.paragraph_format.space_before=Pt(8); pp.paragraph_format.space_after=Pt(8)

def B(t):
    pp=doc.add_paragraph(style='List Bullet'); pp.clear()
    r=pp.add_run(t); r.font.size=Pt(10.5); r.font.name='微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')

def P(t,bold=False):
    pp=doc.add_paragraph(); pp.paragraph_format.first_line_indent=Cm(0.74)
    r=pp.add_run(t); r.font.size=Pt(10.5); r.font.name='微软雅黑'; r.bold=bold
    r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')

def RP(segs):
    pp=doc.add_paragraph(); pp.paragraph_format.first_line_indent=Cm(0.74)
    for txt,bold in segs:
        r=pp.add_run(txt); r.font.size=Pt(10.5); r.font.name='微软雅黑'; r.bold=bold
        r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')

PB=doc.add_page_break
H1=lambda t: doc.add_heading(t,level=1)
H2=lambda t: doc.add_heading(t,level=2)

# 封面
for _ in range(6): doc.add_paragraph()
for txt,sz,col in [('AI一键精修',32,RGBColor(0x1A,0x3C,0x6E)),('市场需求文档（MRD）',22,RGBColor(0x2C,0x5F,0xA0))]:
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=pp.add_run(txt); r.font.size=Pt(sz); r.font.bold=True; r.font.color.rgb=col; r.font.name='微软雅黑'
    r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=pp.add_run('Market Requirements Document'); r.font.size=Pt(14); r.font.color.rgb=RGBColor(0x88,0x88,0x88); r.font.italic=True
for _ in range(4): doc.add_paragraph()
it=doc.add_table(rows=5,cols=2); it.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate([('产品名称','AI一键精修'),('文档版本','V1.0'),('编写日期','2026年8月14日'),('文档状态','初稿'),('密级','内部保密')]):
    for j,val in enumerate([k,v]):
        c=it.rows[i].cells[j]; c.text=''; r=c.paragraphs[0].add_run(val)
        r.font.size=Pt(11); r.font.name='微软雅黑'; r.bold=(j==0); r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
        if j==0: r.font.color.rgb=RGBColor(0x2C,0x5F,0xA0)
PB()

# 目录
H1('目  录')
for item in ['一、市场现状与行业格局','二、用户画像与核心场景','三、市场验证（抖音实测）','四、竞品分析','五、产品差异化优势','六、产品功能规划','七、商业变现','八、后续规划','九、风险与应对','附录']:
    pp=doc.add_paragraph(); r=pp.add_run(item); r.font.size=Pt(12); r.font.name='微软雅黑'
    r.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); pp.paragraph_format.space_after=Pt(8)
PB()

# 第一章
H1('一、市场现状与行业格局')
H2('1.1 修图产品两大分类')
P('当前修图类产品主要分为两大类：')
RP([('第一类：手动修图工具。',True),('以醒图、美图秀秀为代表，用户需要自行操作完成修图。这类产品虽然功能强大，但要求用户具备一定的修图技能和审美能力。',False)])
RP([('第二类：代修服务。',True),('由他人替用户完成修图，又可细分为两种：',False)])
B('人工代修：以闲鱼平台卖家为代表，由真人修图师根据用户需求完成精修；')
B('AI代修：以人工智能技术驱动，用户提交图片和需求后由AI自动完成修图——这正是我们产品的方向。')
IMG('图1-1 修图产品市场分类图')
H2('1.2 手动修图工具的核心痛点')
P('尽管手机端修图App已尽最大程度进行简化，但实际使用中用户仍面临诸多问题：')
RP([('1. 使用门槛高。',True),('用户进入操作后，需要面对图片格式、样式、模板等多种选择。大量模板摆在面前，用户不知道哪个模板好看，只能不断切换尝试。',False)])
RP([('2. 等待时间长。',True),('每切换一个模板样式，需要等待将近10至20秒才能看到效果。频繁切换导致用户等待时间大幅增加，流失率显著升高。',False)])
RP([('3. 积分消耗。',True),('用户在切换模板的过程中，每次操作都会消耗积分。反复试错不仅消耗时间，也直接消耗用户的付费资源。',False)])
RP([('4. 微调成本高。',True),('即使用户选定了一个基本满意的模板，往往还需要进行瘦脸、瘦腿等手动微调，这进一步拉长了使用时间，提升了操作成本。',False)])
H2('1.3 人工代修市场')
P('在手动修图工具体验不佳的背景下，人工代修市场应运而生并持续繁荣：')
B('平台表现：闲鱼平台上代修图服务热销，热门账号销量普遍在1万至2万以上；')
B('价格区间：普通修图单价在2至5元/张，根据精修层次和图片效果，价格最高可达20元/张；')
B('驱动因素：旅游拍照、朋友约会等场景下，用户身体疲劳，不愿花时间自行修图，愿意付费让他人代劳。')
P('这一市场的存在本身就证明了"代修"需求的真实性和付费意愿。')
IMG('图1-2 闲鱼代修服务销量及价格截图示例')
H2('1.4 AI代修的成本优势')
P('AI代修相比人工代修具有显著的成本优势：')
B('技术原理：用户将提示词和需求告诉AI，由AI生成修图结果；')
B('生成成本：1.5K以下图片生成成本约为0.3元/次，2K以上图片约为0.6元/次；')
B('成本对比：人工代修单价为2至20元/张，AI代修成本仅为0.3至0.6元/次，成本差距达到数倍至数十倍。')
P('这一成本优势意味着AI代修可以在保证效果的前提下，以远低于人工的价格向用户提供服务，具备巨大的商业空间。')
TBL(['修图方式','单次成本','耗时','用户操作'],[['手动修图工具','积分消耗','10-30分钟','全程手动操作'],['人工代修','2-20元/张','数小时至数天','提交需求后等待'],['AI代修（我们）','0.3-0.6元/次','数十秒','丢图即可，无需描述']])
PB()

# 第二章
H1('二、用户画像与核心场景')
H2('2.1 精准用户画像')
P('通过抖音市场实测，我们发现产品的核心用户并非通常认知中的"爱拍照、会修图"的女性群体，而是有显著差异的另一类人群。')
P('非目标用户：',bold=True)
P('长相出众、天天拍照、擅长修图的女性用户。这类用户本身基础条件好，使用修图工具时只需调整色温、色调，微调局部即可出片，对工具的依赖程度低，也不需要代修服务。')
P('核心目标用户：',bold=True)
P('长相平平、不会修图的女性用户。这类用户的特点是：')
B('需要调整的部位多，不像长相出众的用户只需微调；')
B('缺乏修图技能，调来调去容易把图片改得四不像；')
B('使用修图工具的时间成本极高，往往折腾很久仍不满意；')
B('有强烈的变美意愿，但缺乏实现手段。')
IMG('图2-1 用户画像对比图（非目标用户 vs 核心目标用户）')
H2('2.2 核心使用场景')
P('目标用户的典型使用场景集中在以下情形：')
B('旅游拍照后：旅途中拍摄大量照片，身体疲惫，回来后没有精力逐张精修；')
B('与朋友约会/出去玩后：活动结束后身心放松，不愿再花时间修图；')
B('其他特殊场合：聚会、活动等需要出片但无暇修图的场景。')
P('这些场景的共性是：用户身体疲劳，不愿花时间自己修图，倾向于寻求代修服务。正是这类场景催生了闲鱼上代修卖家的生意，也正是我们产品的切入点。')
H2('2.3 市场逻辑链')
P('综合以上分析，市场需求的逻辑链条非常清晰：')
P('使用场景（累/不会P）→ 用户痛点（工具难用、自己修不好）→ 现有方案（闲鱼代P，2-20元/张，价格高、等待久）→ 我们的机会（AI代P，速度快、成本低、效果好）',bold=True)
P('这一逻辑链从真实场景出发，经过痛点验证和方案对比，最终指向AI一键精修的产品机会。')
PB()

# 第三章
H1('三、市场验证（抖音实测）')
P('为了验证产品方向和用户需求，我们基于AI技术开发了一个小程序Demo，并在抖音平台进行了真实的市场测试。')
H2('3.1 验证方式')
P('第一阶段：作品发布+投流。',bold=True)
P('在抖音发布修图作品并进行广告投流，以"免费帮用户P图"为吸引点，引导用户提交图片进行体验。')
P('第二阶段：评论区自然获客。',bold=True)
P('停止付费投流后，转为在抖音热门视频和帖子的评论区寻找有修图需求的用户。将用户发到评论中的图片下载下来，用AI精修后上传作品。全程不在评论区进行引流、报价或引导私聊，仅通过展示修图作品本身吸引用户。')
IMG('图3-1 抖音修图作品展示及用户评论截图')
H2('3.2 验证数据')
P('经过一段时间的运营，市场验证取得了积极成果：')
B('服务用户数：累计服务20余位用户；')
B('满意度：90%以上的用户认为修图效果惊艳，甚至以为是人工P的；')
B('AI痕迹问题：个别图片在AI处理过程中痕迹较明显，用户能一眼看出是AI修图，但这类情况占少数；')
B('自然获客：在不做任何引流的情况下，仅通过在评论区展示作品，就吸引了十多位用户主动联系；')
B('作品认可：几乎每条作品的原图作者都会点赞，认为修图效果不错且态度敬业；')
B('合作意向：有2位客户对作品非常喜欢，主动表达了合作意向。')
H2('3.3 定价测试')
P('在市场验证过程中，我们也对定价进行了初步测试：')
B('测试价格：0.5元/次；')
B('使用体验：一张图基本一次就能搞定，效果好、速度快；')
B('用户反馈：用户对价格和效果都表示认可，接受度高。')
P('0.5元/次的定价远低于闲鱼人工代修2元起步的价格，同时仍能保持合理的利润空间，验证了商业模式的可行性。')
IMG('图3-2 用户正向反馈聊天截图')
PB()

# 第四章
H1('四、竞品分析')
H2('4.1 头部App竞品')
P('修图市场已形成两大巨头格局，用户