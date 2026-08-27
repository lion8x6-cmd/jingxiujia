#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将生成的图表插入 MRD docx 文档"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DOCX = '/Users/mr.baihe/WorkBuddy/P 图修图-精修家/output/mrd-20260814/stage3/AI一键精修_MRD_V1.0.docx'
IMG_DIR = '/Users/mr.baihe/WorkBuddy/P 图修图-精修家/output/mrd-20260814/stage2/images'

IMG_MAP = {
    '图1-1 修图产品市场分类图': 'fig1-1-market-classification.png',
    '图2-1 用户画像对比图（非目标用户 vs 核心目标用户）': 'fig2-1-user-persona.png',
    '图5-1 产品使用流程对比图（竞品多步操作 vs 我们三步流程）': 'fig5-1-flow-comparison.png',
    '图5-2 框选局部修改功能示意图': 'fig5-2-selection-edit.png',
    '图5-3 参考图学习功能示意图（待修图 + 参考图 → 效果图）': 'fig5-3-reference-learning.png',
    '图6-1 产品功能架构图': 'fig6-1-feature-architecture.png',
    '图7-1 iOS激活码充值流程示意图': 'fig7-1-ios-payment-flow.png',
    '图8-1 用户心智占位示意图': 'fig8-1-mind-positioning.png',
    '图9-1 SWOT分析矩阵': 'fig9-1-swot.png',
}

doc = Document(DOCX)
replaced = 0

for para in doc.paragraphs:
    text = para.text.strip()
    if text.startswith('\u3010') and text.endswith('\u3011'):
        caption = text[1:-1]
        if caption not in IMG_MAP:
            continue
        img_path = os.path.join(IMG_DIR, IMG_MAP[caption])
        if not os.path.exists(img_path):
            print(f'  NOT FOUND: {img_path}')
            continue

        # 清除占位符
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = para._p.get_or_add_pPr()
        for tag in ['w:pBdr', 'w:shd']:
            for el in pPr.findall(qn(tag)):
                pPr.remove(el)

        # 插入图片
        run = para.add_run()
        run.add_picture(img_path, width=Inches(5.8))

        # 在图片段落后插入图注段落
        new_p = OxmlElement('w:p')
        para._p.addnext(new_p)
        from docx.text.paragraph import Paragraph
        cap_para = Paragraph(new_p, para._parent)
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = cap_para.add_run(caption)
        crun.font.size = Pt(9)
        crun.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        crun.font.name = '微软雅黑'
        crun.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        cap_para.paragraph_format.space_after = Pt(12)
        cap_para.paragraph_format.space_before = Pt(4)

        replaced += 1
        print(f'  OK: {caption}')

print(f'\n共替换 {replaced} 张图片')
doc.save(DOCX)
print(f'已保存: {DOCX}')
