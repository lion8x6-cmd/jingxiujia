#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""MRD 文档渲染引擎"""
import sys, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()
s = doc.styles['Normal']
s.font.name='微软雅黑'; s.font.size=Pt(10.5); s.font.color.rgb=RGBColor(0x33,0x33,0x33)
s.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.5
for sec in doc.sections:
    sec.top_margin=Cm(2.54); sec.bottom_margin=Cm(2.54); sec.left_margin=Cm(3.17); sec.right_margin=Cm(3.17)
for lv,sz in [(1,18),(2,14),(3,12)]:
    h=doc.styles[f'Heading {lv}']; h.font.name='微软雅黑'; h.font.size=Pt(sz); h.font.bold=True
    h.font.color.rgb=[None,RGBColor(0x1A,0x3C,0x6E),RGBColor(0x2C,0x5F,0xA0),RGBColor(0x3A,0x7A,0xC4)][lv]
    h.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    h.paragraph_format.space_before=Pt(18 if lv==1 else 12); h.paragraph_format.space_after=Pt(8)

def _cell_shade(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def _set_font(run, size=10.5, bold=False, color=None, italic=False):
    run.font.name='微软雅黑'; run.font.size=Pt(size); run.bold=bold; run.font.italic=italic
    run.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    if color: run.font.color.rgb=color

def render(items):
    for item in items:
        op = item[0]
        if op == 'h1': doc.add_heading(item[1], level=1)
        elif op == 'h2': doc.add_heading(item[1], level=2)
        elif op == 'h3': doc.add_heading(item[1], level=3)
        elif op == 'pb': doc.add_page_break()
        elif op == 'p':
            pp=doc.add_paragraph(); pp.paragraph_format.first_line_indent=Cm(0.74)
            _set_font(pp.add_run(item[1]), bold=item[2] if len(item)>2 else False)
        elif op == 'rp':
            pp=doc.add_paragraph(); pp.paragraph_format.first_line_indent=Cm(0.74)
            for txt,bold in item[1]: _set_font(pp.add_run(txt), bold=bold)
        elif op == 'b':
            pp=doc.add_paragraph(style='List Bullet'); pp.clear()
            _set_font(pp.add_run(item[1]))
        elif op == 'img':
            pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=pp.add_run(f'\u3010{item[1]}\u3011')
            _set_font(r, size=10, color=RGBColor(0x99,0x99,0x99), italic=True)
            pPr=pp._p.get_or_add_pPr()
            pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/><w:left w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/><w:bottom w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/><w:right w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/></w:pBdr>'))
            pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>'))
            pp.paragraph_format.space_before=Pt(8); pp.paragraph_format.space_after=Pt(8)
        elif op == 'tbl':
            headers, rows = item[1], item[2]
            t=doc.add_table(rows=1+len(rows), cols=len(headers))
            t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
            for i,h in enumerate(headers):
                c=t.rows[0].cells[i]; c.text=''; pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                _set_font(pp.add_run(h), size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
                _cell_shade(c,'2C5FA0')
            for ri,row in enumerate(rows):
                for ci,val in enumerate(row):
                    c=t.rows[ri+1].cells[ci]; c.text=''; pp=c.paragraphs[0]
                    _set_font(pp.add_run(str(val)), size=9.5)
                    if ri%2==1: _cell_shade(c,'E8F0FE')
            doc.add_paragraph()
        elif op == 'cover':
            for _ in range(6): doc.add_paragraph()
            for txt,sz,col in [('AI一键精修',32,RGBColor(0x1A,0x3C,0x6E)),('市场需求文档（MRD）',22,RGBColor(0x2C,0x5F,0xA0))]:
                pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                _set_font(pp.add_run(txt), size=sz, bold=True, color=col)
            pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=pp.add_run('Market Requirements Document'); _set_font(r, size=14, color=RGBColor(0x88,0x88,0x88), italic=True)
            for _ in range(4): doc.add_paragraph()
            it=doc.add_table(rows=5,cols=2); it.alignment=WD_TABLE_ALIGNMENT.CENTER
            for i,(k,v) in enumerate([('产品名称','AI一键精修'),('文档版本','V1.0'),('编写日期','2026年8月14日'),('文档状态','初稿'),('密级','内部保密')]):
                for j,val in enumerate([k,v]):
                    c=it.rows[i].cells[j]; c.text=''; pp=c.paragraphs[0]
                    _set_font(pp.add_run(val), size=11, bold=(j==0), color=RGBColor(0x2C,0x5F,0xA0) if j==0 else None)
        elif op == 'toc':
            doc.add_heading('目  录', level=1)
            for txt in ['一、市场现状与行业格局','二、用户画像与核心场景','三、市场验证（抖音实测）','四、竞品分析','五、产品差异化优势','六、产品功能规划','七、商业变现','八、后续规划','九、风险与应对','附录']:
                pp=doc.add_paragraph(); _set_font(pp.add_run(txt), size=12); pp.paragraph_format.space_after=Pt(8)
        elif op == 'space':
            for _ in range(item[1]): doc.add_paragraph()

if __name__ == '__main__':
    data = json.load(open(sys.argv[1], encoding='utf-8'))
    render(data)
    out = sys.argv[2]
    doc.save(out)
    print(f'SAVED: {out}')
